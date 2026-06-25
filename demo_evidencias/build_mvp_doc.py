from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/alfonsocr/Documents/programacion/uni/EducationMaxxingMVP")
OUT = ROOT / "documentacion_funcionamiento_mvp.docx"
EVIDENCE = ROOT / "demo_evidencias"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, size=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    style_run(run, bold=True, size=16 if level == 1 else 13, color="2E74B5")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    style_run(run, size=11, color="000000")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_scope_table(doc):
    headers = ["FUNCIONALIDAD", "¿INCLUIDA EN EL MVP?", "OBSERVACIÓN"]
    rows = [
        ["Inicio de sesión y autenticación por JWT", "Sí - incluida", "Funcional y probada con administrador, profesor y alumno durante el recorrido del demo."],
        ["Gestión de usuarios por administrador", "Sí - incluida", "El administrador crea usuarios con rol profesor y alumno desde la interfaz web."],
        ["Creación y asignación de cursos", "Sí - incluida", "El administrador crea un curso nuevo y lo asigna al profesor creado en el flujo."],
        ["Inscripción de alumnos a cursos", "Sí - incluida", "El alumno creado se inscribe al curso mediante el panel de administración."],
        ["Creación de asignaciones por profesor", "Sí - incluida", "El profesor entra a su curso y crea una asignación con fecha de inicio y cierre."],
        ["Subida de tareas por alumno", "Sí - incluida", "El alumno adjunta un archivo válido y recibe confirmación de entrega."],
        ["Validación de archivo inválido", "Sí - incluida", "El sistema rechaza un ejecutable y muestra el mensaje de error correspondiente."],
        ["Revisión y calificación de entregas", "Sí - parcial", "El backend y la vista de profesor contemplan revisión; no fue el foco principal de este demo."],
        ["Reportes académicos avanzados", "No - V2", "Quedan fuera del alcance funcional demostrado en esta sección del MVP."],
        ["Persistencia en base de datos real", "No - V2", "El MVP usa repositorios en memoria y almacenamiento local de archivos para evidenciar el flujo."],
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    widths = [Inches(2.15), Inches(1.45), Inches(2.75)]
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        style_run(run, bold=True, size=9.5, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = widths[idx]
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            style_run(run, size=9)
    doc.add_paragraph()


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    style_run(run, bold=True, size=9, color="555555")
    return p


def add_image_evidence(doc, title, image_name, description):
    add_heading(doc, title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(EVIDENCE / image_name), width=Inches(6.15))
    add_caption(doc, f"Evidencia: {image_name}")
    add_body(doc, description)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("6. Funcionamiento del MVP")
    style_run(run, bold=True, size=22, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("EducationMaxxingMVP - Evidencia funcional del demo")
    style_run(run, size=11, color="555555")

    add_body(
        doc,
        "Esta sección documenta el alcance funcional del demo, los pasos necesarios para ejecutar el sistema y las evidencias capturadas durante un recorrido completo por la interfaz web. El flujo validado cubre administración de usuarios, creación de curso, creación de asignación y entrega de tarea por parte de un alumno.",
    )

    add_heading(doc, "6.1 Alcance Funcional del Demo", level=1)
    add_body(
        doc,
        "El demo se limita al comportamiento central del MVP: autenticación por roles, administración básica de usuarios y cursos, creación de asignaciones y carga de archivos de entrega. Las funciones marcadas como V2 representan capacidades fuera del alcance inmediato o no demostradas dentro del recorrido principal.",
    )
    add_scope_table(doc)

    add_heading(doc, "6.2 Instrucciones de Ejecución del Demo", level=1)
    add_heading(doc, "Requisitos previos", level=2)
    for item in [
        "Go 1.26.2 o compatible para ejecutar el backend.",
        "Node.js 25.9.0 y npm 11.12.1 o versiones compatibles con Next.js 16.2.7.",
        "Navegador moderno, preferentemente Google Chrome.",
        "Puerto 8080 disponible para la API y puerto 3000 disponible para el frontend.",
        "Archivo frontend/.env.local con NEXT_PUBLIC_API_URL=http://localhost:8080/api.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Pasos de instalación y ejecución", level=2)
    for step in [
        "Abrir una terminal en la raíz del proyecto EducationMaxxingMVP.",
        "Ejecutar go mod download para descargar las dependencias del backend.",
        "Iniciar el backend con go run ./cmd. La API quedará disponible en http://localhost:8080.",
        "Abrir una segunda terminal y entrar a la carpeta frontend.",
        "Ejecutar npm install si las dependencias del frontend aún no están instaladas.",
        "Verificar que frontend/.env.local apunte a http://localhost:8080/api.",
        "Iniciar la interfaz con npm run dev. El proyecto está configurado para usar Webpack y evitar el fallo previo de Turbopack.",
        "Abrir http://localhost:3000/login en el navegador.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "Flujo básico de navegación", level=2)
    for item in [
        "Ingresar como administrador con admin@educationmaxxing.com y admin1234.",
        "Desde Users, crear un profesor y un alumno. Después, desde Courses, crear un curso y asignarlo al profesor.",
        "Regresar a Users e inscribir el alumno en el curso creado.",
        "Cerrar sesión e ingresar como profesor para abrir el curso y crear una asignación.",
        "Cerrar sesión e ingresar como alumno para abrir la asignación, probar un archivo inválido y posteriormente subir un archivo válido.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Datos usados en la captura", level=2)
    add_body(doc, "Profesor: profesor.demo.1782352993821@university.edu. Alumno: alumno.demo.1782352993821@university.edu. Curso: Arquitectura MVP 93821. Asignación: Entrega Demo 93821.")

    add_heading(doc, "6.3 Evidencias del Sistema en Ejecución", level=1)
    add_image_evidence(
        doc,
        "Captura 1: Pantalla de inicio del administrador",
        "01_inicio_dashboard_admin.png",
        "La captura muestra el ingreso exitoso del usuario administrador al panel principal del sistema. En esta pantalla se identifica la capa de presentación web, compuesta por el layout autenticado, el menú lateral y la vista de resumen de plataforma. Arquitectónicamente, este estado confirma que el login fue validado por el handler HTTP de autenticación, que el token JWT fue almacenado por el frontend y que la sesión permite cargar vistas protegidas.",
    )
    add_image_evidence(
        doc,
        "Captura 2: Entrega exitosa de una tarea",
        "02_entrega_exitosa.png",
        "La captura muestra el caso de uso principal del alumno: subir un archivo válido a una asignación existente. El sistema devuelve una confirmación con identificador de entrega y registra el historial de la submission como confirmada. En la arquitectura se observa la interacción entre el handler de entregas, SubmissionService, el repositorio de submissions y el puerto FileStorage, que guarda físicamente el archivo en almacenamiento local.",
    )
    add_image_evidence(
        doc,
        "Captura 3: Manejo de error por archivo inválido",
        "03_error_archivo_invalido.png",
        "La captura evidencia el manejo de excepciones cuando el alumno intenta subir un archivo ejecutable no permitido. La interfaz conserva el contexto de la asignación y muestra el mensaje file type is not allowed sin romper el flujo del usuario. Este comportamiento está alineado con la capa de aplicación, donde SubmissionService valida el tipo de archivo antes de persistirlo, y con la capa HTTP, que transforma el error en una respuesta controlada.",
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("EducationMaxxingMVP - Documentación del demo")
        style_run(run, size=9, color="555555")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
